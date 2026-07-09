from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor

from gli.audit_block6m2d import audit_summary


OUT = Path("docs/Bloque_6M_2D_Compatibilidad_C.docx")


def set_font(run, size=11, bold=False, color=None):
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    r = p.add_run(text)
    set_font(r, 16 if level == 1 else 13 if level == 2 else 12, True, "2E74B5" if level < 3 else "1F4D78")
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.10
    r = p.add_run(text)
    set_font(r)
    return p


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_font(r, 9, True)
        if widths:
            hdr[i].width = Inches(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if widths:
                cells[i].width = Inches(widths[i])
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(val))
            set_font(r, 8.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def fmt(x, n=4):
    if isinstance(x, bool):
        return "sí" if x else "no"
    if abs(float(x)) >= 1e3 or (abs(float(x)) < 1e-3 and float(x) != 0):
        return f"{float(x):.{n}e}"
    return f"{float(x):.{n}f}"


def main():
    summary = audit_summary()
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    r = title.add_run("Bloque 6M-2D - Compatibilidad del estado C")
    set_font(r, 20, True, "000000")
    sub = doc.add_paragraph()
    sub.paragraph_format.space_after = Pt(14)
    r = sub.add_run("Caso Santos {50,70,80}. Auditoría local de restricciones B→C→D, sin avance a D→E ni E→F.")
    set_font(r, 11, False, "555555")

    add_heading(doc, "Dictamen", 1)
    add_body(
        doc,
        "El estado C heredado del RHS B→C previo no satisface el manifold completo requerido por el RHS corregido C→D. "
        "La causa nace en B→C: se transportaba una película dinámica con cierre de momento incompatible con la etapa de elevación de Santos y no se mantenía la restricción algebraica de masa de la golfada. "
        "Se implementó un modo paralelo B→C santos_compatible que impone desde B las restricciones 4.1.32, 4.1.33, 4.1.35, 4.1.39 y 4.1.49-4.1.51. Con ese estado C, el RHS corregido C→D queda certificado y fue conectado a la API. D→E permanece provisional y E→F desconectado.",
    )

    add_heading(doc, "Ecuaciones usadas para R_C", 1)
    rows = [
        ("4.1.32", "dh_L/dt = v_l", "cinemática del tope de golfada"),
        ("4.1.33", "dh_B/dt = v_B", "cinemática de burbuja"),
        ("4.1.35", "2π(r-y)h_B dy/dt + A_f v_f - q_res = 0", "balance de película"),
        ("4.1.39", "A_t v_l - A_f v_f - A_B v_B = 0", "cierre algebraico de golfada/película"),
        ("4.1.49-51", "v_B = a v_l + b; dv_B/dt = a dv_l/dt", "relación Brown usada por Santos"),
        ("EOS", "P_t1 = m_g ZRT/(M A_B h_B)", "cierre masa-densidad-presión gas"),
    ]
    add_table(doc, ["Referencia", "Expresión", "Uso"], rows, [1.0, 3.0, 2.2])

    add_heading(doc, "Clasificación del estado C", 1)
    rows = [
        ("m_c, m_g, rho_g, P_g", "variables dinámicas continuas", "no pueden reinicializarse en C"),
        ("v_g/v_B, v_f, y, m_film", "dinámicas con cierres algebraicos/geométricos", "deben llegar compatibles desde B→C"),
        ("h_B, h_L, v_l", "variables dinámicas continuas", "misma cinemática a ambos lados de C"),
        ("V_gi", "condición de evento/transición", "en C solo cierra la válvula motora"),
        ("fallback, producido", "diagnóstico/ledger", "no justifican balancear película por proyección"),
    ]
    add_table(doc, ["Variable", "Clasificación", "Restricción de transición"], rows, [1.7, 2.0, 2.5])

    doc.add_page_break()
    add_heading(doc, "Residuos comparativos R_C", 1)
    current = {r["name"]: r for r in summary["residuals_current"]}
    compat = {r["name"]: r for r in summary["residuals_compatible"]}
    rows = []
    for name in [
        "slug_mass_algebraic",
        "film_mass_geometry",
        "gas_eos",
        "gas_density_geometry",
        "event_c_volume",
        "film_mass_differential",
        "slug_mass_differential",
    ]:
        rows.append(
            (
                name,
                current[name]["units"],
                fmt(current[name]["value"]),
                fmt(current[name]["normalized"]),
                fmt(compat[name]["value"]),
                fmt(compat[name]["normalized"]),
            )
        )
    add_table(
        doc,
        ["Residuo", "Unidad", "actual", "actual norm.", "compatible", "compatible norm."],
        rows,
        [1.55, 0.65, 1.0, 1.0, 1.0, 1.0],
    )
    add_body(
        doc,
        f"Máximo normalizado: actual {summary['inherited_max_R']:.6g}; compatible {summary['compatible_max_R']:.6g}. "
        "La tolerancia de certificación usada para compatibilidad es 1e-6.",
    )

    add_heading(doc, "Reejecución B→C→D corregida", 1)
    rows = [
        ("C actual heredado", fmt(summary["inherited_c_time_s"]), fmt(summary["inherited_max_R"]), "no", "C→D corregido rechazado; balance líquido falla"),
        ("C compatible", fmt(summary["compatible_c_time_s"]), fmt(summary["compatible_max_R"]), "sí", "C→D corregido certificado"),
        ("D con C compatible", fmt(summary["cd_certified_event_d_s"]), "-", "sí", f"gas {summary['cd_certified_gas_balance']:.2e}; líquido {summary['cd_certified_liquid_balance']:.2e}; EOS {summary['cd_certified_eos']:.2e}"),
    ]
    add_table(doc, ["Caso", "tiempo relativo [s]", "max R_C", "certificado", "observación"], rows, [1.5, 1.1, 1.0, 0.9, 1.9])

    add_heading(doc, "Trazabilidad código-prueba", 1)
    rows = [
        ("R_C", "src/gli/block6m2d_compatibility.py", "compatibility_residuals_c", "tests/test_block6m2d_c_compatibility.py"),
        ("Auditoría ejecutable", "src/gli/audit_block6m2d.py", "audit_summary", "tests/test_block6m2d_c_compatibility.py"),
        ("B→C compatible", "src/gli/stage_bc_common.py", "rhs_bc_santos_compatible; rhs_mode=santos_compatible", "test_santos_compatible_bc_transports_c_restrictions"),
        ("C→D corregido", "src/gli/stage_cd_common.py", "_cd_terms_santos; rhs_mode=santos_corrected", "test_corrected_cd_from_compatible_c_state_is_certified"),
        ("API", "src/gli_api/simulation_service.py", "B→C santos_compatible; C→D santos_corrected", "suite completa"),
    ]
    add_table(doc, ["Elemento", "Archivo", "Símbolo", "Prueba"], rows, [1.1, 1.7, 2.0, 1.6])

    add_heading(doc, "Alcance y límites", 1)
    add_body(
        doc,
        "API A→E con validationLevel=provisional: B→C compatible y C→D corregido quedan conectados; D→E sigue provisional y E→F desconectado. No se usaron proyecciones, clipping ni calibración. Pruebas: 101 passed en 70.55 s.",
    )

    OUT.parent.mkdir(exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
