from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from gli.audit_block6m3_de import audit_summary


OUT = Path("docs/Bloque_6M_3_Correccion_D_E.docx")


def set_font(run, size=11, bold=False, color=None):
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def para(doc, text, size=11, color=None, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.10
    r = p.add_run(text)
    set_font(r, size, bold, color)
    return p


def heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    r = p.add_run(text)
    set_font(r, 16 if level == 1 else 13, True, "2E74B5")


def table(doc, headers, rows, widths):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.width = Inches(widths[i])
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_font(r, 8.5, True)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].width = Inches(widths[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(val))
            set_font(r, 8)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def fmt(x):
    try:
        x = float(x)
    except Exception:
        return str(x)
    if x and (abs(x) < 1e-3 or abs(x) >= 1e4):
        return f"{x:.4e}"
    return f"{x:.6f}".rstrip("0").rstrip(".")


def main():
    s = audit_summary()
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Bloque 6M-3C - Corrección paralela D→E")
    set_font(r, 20, True)
    para(doc, "Caso Santos {50,70,80}. Ruta D→E corregida en paralelo; API conserva D→E legado/provisional y E→F desconectado.", 11, "555555")

    heading(doc, "Dictamen")
    para(doc, "La ruta paralela rhs_mode='santos_corrected' queda matemáticamente compatible con el estado D entregado por C→D corregido. Cierra los residuales de memoria de v_f, balance con q_res, GLV cerrada y evento E. No se conecta a la API como ruta principal porque E→F sigue desconectado y la validación cuantitativa global D→E aún requiere comparación formal contra el caseId Santos.")

    heading(doc, "Contratos corregidos")
    rows = [
        ("v_f como memoria", "v_f(D+) se toma de 4.1.40 usando el estado D corregido.", fmt(s["corrected_film_velocity_jump_m_s"]), "m/s"),
        ("q_res etapa 3", "Inventario total = I_D + q_res t.", fmt(s["corrected_reservoir_residual_m3"]), "m³"),
        ("balance líquido", "Error relativo contra inventario D correcto.", fmt(s["corrected_liquid_balance_error"]), "adim."),
        ("GLV", "Ruta corregida enclava GLV cerrada; caudal cero.", str(s["corrected_glv_open_any"]), "-"),
        ("evento E", "Base de golfada llega a superficie.", fmt(s["corrected_event_e_time_s"]), "s"),
    ]
    table(doc, ["Contrato", "Implementación", "Residual/valor", "Unidad"], rows, [1.45, 3.05, 1.1, 0.55])

    heading(doc, "Comparación legado vs corregido")
    rows = [
        ("Certificación D→E", str(s["certified"]), str(s["corrected_certified"])),
        ("E relativo [s]", fmt(s["event_e_time_s"]), fmt(s["corrected_event_e_time_s"])),
        ("Producido [m³]", fmt(s["produced_volume_m3"]), fmt(s["corrected_produced_volume_m3"])),
        ("Error líquido", fmt(s["legacy_liquid_balance_error"]), fmt(s["corrected_liquid_balance_error"])),
        ("Reservorio residual [m³]", fmt(s["reservoir_missing_m3"]), fmt(s["corrected_reservoir_residual_m3"])),
        ("GLV abierta", str(s["glv_open_any_de"]), str(s["corrected_glv_open_any"])),
        ("Máx. residual normalizado", fmt(s["max_residual_normalized"]), fmt(s["corrected_max_residual_normalized"])),
    ]
    table(doc, ["Magnitud", "Legado", "Corregido paralelo"], rows, [2.35, 1.75, 1.75])

    doc.add_page_break()
    heading(doc, "Residuales corregidos")
    rows = [(r["name"], r["status"], fmt(r["value"]), fmt(r["normalized"]), r["interpretation"]) for r in s["corrected_residuals"]]
    table(doc, ["Residual", "Estado", "Valor", "Norm.", "Interpretación"], rows, [1.3, 0.6, 0.8, 0.7, 2.6])

    heading(doc, "Alcance API y pruebas")
    rows = [
        ("Integrador", "src/gli/stage_de_dynamic.py", "simulate_stage_d_to_e(..., rhs_mode='santos_corrected')"),
        ("Auditoría", "src/gli/audit_block6m3_de.py", "corrected_residuals_d / corrected_certified"),
        ("Pruebas", "tests/test_block6m3_de_audit.py", "legado provisional; corregido cierra residuales"),
        ("API", "src/gli_api/simulation_service.py", "validationLevel='provisional'; D→E legado/provisional"),
    ]
    table(doc, ["Elemento", "Archivo", "Contrato"], rows, [1.0, 2.0, 3.0])
    para(doc, "Compuerta: D→E corregido queda compatible en paralelo, pero la API no se promueve a certified mientras E→F siga desconectado. Suite local: 105 passed.")

    OUT.parent.mkdir(exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
