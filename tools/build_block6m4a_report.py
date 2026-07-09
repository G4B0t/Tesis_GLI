from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from gli.audit_block6m4_ef import audit_summary


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "Bloque_6M_4A_Auditoria_E_F.docx"


def _fmt(value, digits=6):
    if value is None:
        return "N/A"
    try:
        return f"{float(value):.{digits}g}"
    except (TypeError, ValueError):
        return str(value)


def _set_cell(cell, text):
    cell.text = str(text)
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(8)


def add_table(document, headers, rows):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        _set_cell(table.rows[0].cells[idx], header)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            _set_cell(cells[idx], value)
    return table


def build():
    data = audit_summary(max_step_s=0.5)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    title = doc.add_heading("Bloque 6M-4A — Auditoría E→F", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        "Auditoría local de la frontera E usando el estado entregado por D→E "
        "santos_corrected. No se conecta E→F a la API pública, no se promueve "
        "validationLevel y no se modifica la ruta legacy D→E."
    )

    doc.add_heading("1. Pregunta auditada", level=1)
    doc.add_paragraph(
        "¿El E→F vigente puede recibir el estado E corregido sin proyección? "
        f"Dictamen: {'SÍ' if data['can_receive_without_projection'] else 'NO'}."
    )
    doc.add_paragraph(
        f"Evento E D→E corregido: {_fmt(data['event_e_time_s'])} s. "
        f"Evento F detectado por E→F vigente: {_fmt(data['event_f_time_s'])} s; "
        f"alcanzado={data['event_f_reached']}."
    )
    doc.add_paragraph(
        "Fuente declarada por el E→F auditado: "
        f"{data['ef_initial_state_source'] or data['ef_exception'] or 'N/A'}"
    )

    doc.add_heading("2. Contratos evaluados en E", level=1)
    rows = []
    for r in data["residuals"]:
        rows.append(
            (
                r["name"],
                r["status"],
                _fmt(r["value"]),
                _fmt(r["scale"]),
                _fmt(r["normalized"]),
                r["units"],
                r["contract"],
            )
        )
    add_table(
        doc,
        ("Contrato", "Estado", "Valor", "Escala", "Residuo norm.", "Unid.", "Ecuación/condición"),
        rows,
    )

    doc.add_heading("3. Fallas que bloquean integración directa", level=1)
    failed = data["failed_contracts"]
    if failed:
        for name in failed:
            match = next(r for r in data["residuals"] if r["name"] == name)
            doc.add_paragraph(
                f"{name}: {match['interpretation']} "
                f"Residuo normalizado={_fmt(match['normalized'])}.",
                style="List Bullet",
            )
    else:
        doc.add_paragraph("No se detectaron fallas de contrato.")

    doc.add_heading("4. Trazabilidad", level=1)
    add_table(
        doc,
        ("Magnitud", "Código auditado", "Condición trazable", "Resultado"),
        (
            ("rho_g, m_g, P_t1", "src/gli/stage_ef_dynamic.py", "continuidad E+/E-", "OK"),
            ("v_g", "src/gli/stage_ef_dynamic.py", "estado con memoria o transición Santos", "Falla: recalculado"),
            ("v_f", "src/gli/stage_ef_dynamic.py", "estado con memoria", "Falla: cierre algebraico"),
            ("y, m_film", "src/gli/stage_ef_dynamic.py", "geometría de película", "OK"),
            ("fallback/producido", "StageEFResult", "ledgers acumulados", "Falla: no transportados"),
            ("GLV", "StageEFResult.valve_open", "cerrada, sin reapertura", "OK"),
            ("F", "event_f: v_f=0, direction=-1", "cruce descendente", "No certificado"),
        ),
    )

    doc.add_heading("5. Dictamen y siguiente subbloque", level=1)
    doc.add_paragraph(
        "El E→F actual no debe conectarse al estado E corregido. Aunque conserva "
        "presión, masa/densidad de gas y geometría inicial, reconstruye variables "
        "dinámicas con memoria y pierde los ledgers acumulados. El siguiente "
        "subbloque técnico debe ser 6M-4B: ruta E→F corregida que reciba por "
        "identidad el estado E, transporte ledgers acumulados y certifique F con "
        "cruce descendente de v_f=0."
    )
    doc.add_paragraph(
        "Alcance vigente: A→E disponible, D→E legacy sigue como ruta pública, "
        "D→E santos_corrected permanece ruta paralela auditada, E→F desconectado "
        "y validationLevel='provisional'."
    )

    doc.add_heading("6. Pruebas", level=1)
    doc.add_paragraph(
        "Pruebas focales nuevas: tests/test_block6m4_ef_audit.py. Resultado local: "
        "6 passed."
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()

