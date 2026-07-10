from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from gli.audit_block6m5_af import audit_summary
from gli_api.schemas import SimulationInputs
from gli_api.simulation_service import build_parameters, simulate


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "Bloque_6M_5_Certificacion_A_F.docx"


def fmt(value, digits=6):
    try:
        return f"{float(value):.{digits}g}"
    except (TypeError, ValueError):
        return str(value)


def cell(cell, text, bold=False):
    cell.text = str(text)
    for p in cell.paragraphs:
        p.paragraph_format.space_after = Pt(2)
        for r in p.runs:
            r.font.size = Pt(8)
            r.bold = bold


def table(doc, headers, rows):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    for i, h in enumerate(headers):
        cell(tbl.rows[0].cells[i], h, True)
    for row in rows:
        cells = tbl.add_row().cells
        for i, value in enumerate(row):
            cell(cells[i], value)


def build():
    inputs = SimulationInputs(
        tubingDiameter=0.050673,
        valveDepth=1480.0,
        slugLength=412.5,
        surfaceTubingPressure=0.788,
        injectionPressure=6.966,
        api=40.0,
        bsw=50.0,
        gasRelativeDensity=0.7,
        casingPressureOpenRatio=0.7,
        projectName="QA",
        projectistName="QA",
    )
    params = build_parameters(inputs)
    data = audit_summary(params, max_step_s=0.5)
    api_resolution = audit_summary(params, max_step_s=None)
    api_result = simulate(inputs)
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.75)
    sec.bottom_margin = Inches(0.75)
    sec.left_margin = Inches(0.75)
    sec.right_margin = Inches(0.75)
    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10.5)
    for name in ("Heading 1", "Heading 2"):
        styles[name].font.color.rgb = RGBColor(46, 116, 181)

    title = doc.add_heading("Bloque 6M-5 — Certificación A→F", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        "Integración controlada de la cadena Santos corregida: B→C santos_compatible, "
        "C→D santos_corrected, D→E santos_corrected y E→F santos_corrected."
    )

    doc.add_heading("1. Dictamen", level=1)
    doc.add_paragraph(
        f"Certificación A→F: {'SÍ' if data['certified'] else 'NO'}. "
        f"validationLevel candidato: {data['validation_level_candidate']}. "
        f"Evento terminal: {data['terminal_event']}."
    )
    doc.add_paragraph(f"Residuo máximo normalizado: {fmt(data['max_residual_normalized'])}.")
    doc.add_paragraph(
        "Nota temporal: con los parámetros exactos del endpoint, la auditoría global con max_step_s=0.5 reporta "
        f"F={fmt(data['event_times_s']['F_FILM_VELOCITY_ZERO'])} s. "
        "La API usa los pasos nativos/default de cada tramo y coincide con la "
        f"auditoría en resolución API: F={fmt(api_resolution['event_times_s']['F_FILM_VELOCITY_ZERO'])} s; "
        f"el último punto público es F={fmt(api_result.points[-1].t)} s. "
        "El valor 526.778 s observado en la auditoría interna previa correspondía al helper santos_50_70_80(), "
        "no a los parámetros construidos por el endpoint público."
    )

    doc.add_heading("2. Eventos y duraciones", level=1)
    table(doc, ("Evento", "t absoluto [s]"), [(k, fmt(v)) for k, v in data["event_times_s"].items()])
    doc.add_paragraph(
        "Para el endpoint público, el tiempo terminal exacto es el de resolución API: "
        f"{fmt(api_resolution['event_times_s']['F_FILM_VELOCITY_ZERO'])} s."
    )
    doc.add_paragraph("")
    table(doc, ("Tramo", "duración [s]"), [(k, fmt(v)) for k, v in data["stage_durations_s"].items()])

    doc.add_heading("3. Compuertas de certificación", level=1)
    table(
        doc,
        ("Contrato", "Estado", "Residual norm.", "Interpretación"),
        [(r["name"], r["status"], fmt(r["normalized"]), r["interpretation"]) for r in data["residuals"]],
    )

    doc.add_heading("4. API y contrato temporal", level=1)
    table(
        doc,
        ("Elemento", "Resultado"),
        (
            ("API pública", "A_TO_F certified solo porque las compuertas A→F cierran."),
            ("terminalEvent", "F_FILM_VELOCITY_ZERO"),
            ("timeline/events", "A_INITIAL_STATE, B, C, D, E, F en orden estricto"),
            ("timeline/segments", "A_B, B_C, C_D, D_E, E_F contiguos"),
            ("validationLevel", "certified para caseId santos-gli-50-70-80"),
        ),
    )

    doc.add_heading("5. Alcance físico declarado", level=1)
    doc.add_paragraph(
        "La certificación es local al caso Santos {50,70,80}, al conjunto de parámetros "
        "explícitos ya registrado y a los cierres documentados. Liao Tabla 5.14 permanece "
        "como benchmark parcial, no como validación cuantitativa del mismo caseId."
    )

    doc.add_heading("6. Pruebas", level=1)
    doc.add_paragraph(
        "Pruebas focales añadidas: tests/test_block6m5_af_certification.py. "
        "Pruebas de API/timeline actualizadas para A→F."
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
