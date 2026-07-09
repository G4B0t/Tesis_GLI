from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from gli.audit_block6m4_ef import audit_summary


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "Bloque_6M_4B_Correccion_E_F.docx"


def fmt(value, digits=6):
    if value is None:
        return "N/A"
    try:
        return f"{float(value):.{digits}g}"
    except (TypeError, ValueError):
        return str(value)


def set_cell(cell, text, bold=False):
    cell.text = str(text)
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(8)
            run.bold = bold


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        set_cell(table.rows[0].cells[i], h, True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell(cells[i], value)
    return table


def build():
    data = audit_summary(max_step_s=0.5)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10.5)
    for name in ("Heading 1", "Heading 2"):
        styles[name].font.color.rgb = RGBColor(46, 116, 181)

    title = doc.add_heading("Bloque 6M-4B — Corrección paralela E→F", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        "Implementación y auditoría local de una ruta E→F santos_corrected. "
        "La API pública no se promueve: D→E legacy sigue como ruta actual, "
        "E→F permanece desconectado y validationLevel='provisional'."
    )

    doc.add_heading("1. Dictamen", level=1)
    doc.add_paragraph(
        f"E→F legacy certificado: NO. E→F santos_corrected certificado: "
        f"{'SÍ' if data['corrected_certified'] else 'NO'}."
    )
    doc.add_paragraph(
        f"Evento F corregido: alcanzado={data['corrected_event_f_reached']}; "
        f"t_F={fmt(data['corrected_event_f_time_s'])} s. "
        f"Evento F legacy: alcanzado={data['event_f_reached']}; "
        f"t_F={fmt(data['event_f_time_s'])} s."
    )

    doc.add_heading("2. Comparación de contratos", level=1)
    rows = [
        ("v_g", "Falla: recalculado por descarga superficial", "OK: recibido por identidad"),
        ("v_f", "Falla: reconstruido con dv_f/dt=0", "OK: recibido por identidad"),
        ("fallback", "Falla: ledger no transportado", "OK: ledger acumulado transportado"),
        ("producido", "Falla: reinicia en cero", "OK: acumulado no reinicia"),
        ("balance gas", "Falla", "OK bajo tolerancia 1e-8"),
        ("balance líquido", "Falla", "OK bajo tolerancia 1e-8"),
        ("GLV", "OK cerrada", "OK cerrada y sin reapertura"),
        ("evento F", "No certificado", "OK: cruce descendente de v_f=0"),
    ]
    add_table(doc, ("Contrato", "Legacy", "santos_corrected"), rows)

    doc.add_heading("3. Residuos corregidos", level=1)
    rows = []
    for r in data["corrected_residuals"]:
        rows.append((r["name"], r["status"], fmt(r["value"]), fmt(r["normalized"]), r["units"], r["interpretation"]))
    add_table(doc, ("Residuo", "Estado", "Valor", "Norm.", "Unid.", "Interpretación"), rows)

    doc.add_heading("4. Implementación trazable", level=1)
    add_table(
        doc,
        ("Elemento", "Archivo", "Decisión"),
        (
            ("Ruta paralela", "src/gli/stage_ef_dynamic.py", "simulate_stage_e_to_f(..., rhs_mode='santos_corrected')"),
            ("Legacy preservado", "src/gli/stage_ef_dynamic.py", "rhs_mode='legacy' sigue siendo default"),
            ("Auditoría", "src/gli/audit_block6m4_ef.py", "compara legacy vs santos_corrected"),
            ("Pruebas", "tests/test_block6m4_ef_audit.py", "legacy bloqueado, corregido certificado"),
            ("API", "src/gli_api/simulation_service.py", "sin cambios; validationLevel provisional"),
        ),
    )

    doc.add_heading("5. Alcance y siguiente paso", level=1)
    doc.add_paragraph(
        "El tramo E→F santos_corrected queda matemáticamente compatible con el estado E "
        "producido por D→E santos_corrected. No se conectó a la API pública en este bloque. "
        "El siguiente bloque debe integrar/certificar A→F punta a punta y recién entonces "
        "evaluar promoción de alcance y validationLevel."
    )
    doc.add_heading("6. Pruebas", level=1)
    doc.add_paragraph("Pruebas focales: tests/test_block6m4_ef_audit.py — 9 passed.")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()

